"""Generate the static web app data file from the official roster.

Usage:
  python scripts/build_data.py path/to/roster.docx
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

import openpyxl
from docx import Document


MAIN_SHEET = "應變小組主表"
ANNEXES = {
    "附表一_高中專任": {
        "group": "搶救組", "fire_group": "滅火班", "name_index": 1, "title_index": 2
    },
    "附表二_國高中導師": {
        "group": "避難引導組", "fire_group": "避難引導班", "name_index": 2, "title_index": 1
    },
    "附表三_國中專任": {
        "group": "安全防護組", "fire_group": "安全防護班", "name_index": 1, "title_index": 2
    },
}
PLACEHOLDER_NAMES = {"高中專任（含外師）", "國高中導師", "國中專任"}
DOCX_ANNEXES = {
    2: {"group": "搶救組", "fire_group": "滅火班", "source": "附表一_高中專任"},
    3: {"group": "避難引導組", "fire_group": "避難引導班", "source": "附表二_國高中導師"},
    4: {"group": "安全防護組", "fire_group": "安全防護班", "source": "附表三_國中專任"},
}


def text(value: object, fallback: str = "-") -> str:
    value = str(value or "").strip()
    return value or fallback


def normalize_role(value: object) -> str:
    role = text(value, "組員")
    if "組長" in role:
        return "組長"
    if "組員" in role:
        return "組員"
    return role


def keep_first_assignment_per_person(
    records: list[dict[str, str]],
) -> list[dict[str, str]]:
    """Keep one assignment per person, preserving source and row priority.

    Main-sheet rows are read before annex rows, so an explicit main-sheet
    assignment takes precedence over a generic annex assignment. If a person
    appears more than once in the same source sequence, the first row wins.
    """
    output: list[dict[str, str]] = []
    seen_names: set[str] = set()
    for record in records:
        name = record["name"]
        if name in seen_names:
            continue
        seen_names.add(name)
        output.append(record)
    return output


def read_xlsx_records(path: Path) -> list[dict[str, str]]:
    workbook = openpyxl.load_workbook(path, data_only=True, read_only=True)
    records: list[dict[str, str]] = []

    main = workbook[MAIN_SHEET]
    for row in main.iter_rows(min_row=3, max_col=7, values_only=True):
        _, group, fire_group, role, name, title, detail = row
        name = text(name, "")
        if not name or name in PLACEHOLDER_NAMES:
            continue
        records.append(
            {
                "name": name,
                "title": text(title),
                "group": text(group),
                "fireGroup": text(fire_group),
                "role": text(role),
                "detail": text(detail),
                "source": MAIN_SHEET,
            }
        )

    for sheet_name, config in ANNEXES.items():
        sheet = workbook[sheet_name]
        for row in sheet.iter_rows(min_row=4, max_col=5, values_only=True):
            name = row[config["name_index"]]
            title = row[config["title_index"]]
            assignment = row[3]
            detail = row[4]
            name = text(name, "")
            if not name:
                continue
            records.append(
                {
                    "name": name,
                    "title": text(title),
                    "group": config["group"],
                    "fireGroup": config["fire_group"],
                    "role": normalize_role(assignment),
                    "detail": text(detail),
                    "source": sheet_name,
                }
            )

    return keep_first_assignment_per_person(records)


def compact(value: object) -> str:
    return re.sub(r"\s+", "", str(value or ""))


def cell_items(cell) -> list[str]:
    return [re.sub(r"\s+", " ", paragraph.text).strip() for paragraph in cell.paragraphs if paragraph.text.strip()]


def parse_docx_group(value: str) -> tuple[str, str]:
    value = compact(value)
    if value in {"指揮官", "指揮官代理人", "發言人"}:
        return "指揮部", "-"
    match = re.fullmatch(r"(.+?)（(.+?)）", value)
    return (match.group(1), match.group(2)) if match else (value, "-")


def read_docx_records(path: Path) -> list[dict[str, str]]:
    """Read the latest Word roster, keeping its first assignment per person."""
    document = Document(path)
    records: list[dict[str, str]] = []
    group_details: dict[str, str] = {}

    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            if row_index == 1 or len(row.cells) < 5:
                continue
            names = cell_items(row.cells[2])
            titles = cell_items(row.cells[3])
            if len(names) == 1 and len(titles) > 1:
                titles = ["".join(titles)]
            if len(names) != len(titles):
                raise ValueError(f"姓名與職稱未對齊：表 {table_index}、列 {row_index}")

            if table_index in DOCX_ANNEXES:
                annex = DOCX_ANNEXES[table_index]
                group, fire_group, source = annex["group"], annex["fire_group"], annex["source"]
                detail = group_details[group]
            else:
                group, fire_group = parse_docx_group(row.cells[0].text)
                source = MAIN_SHEET
                detail = "\n".join(p.text.strip() for p in row.cells[4].paragraphs if p.text.strip())
                if group != "指揮部":
                    group_details[group] = detail

            role = compact(row.cells[1].text)
            for name, title in zip(names, titles):
                if name.startswith("*"):
                    continue
                records.append(
                    {
                        "name": name,
                        "title": title,
                        "group": group,
                        "fireGroup": fire_group,
                        "role": role,
                        "detail": detail,
                        "source": source,
                    }
                )

    return keep_first_assignment_per_person(records)


def read_records(path: Path) -> list[dict[str, str]]:
    if path.suffix.lower() == ".docx":
        return read_docx_records(path)
    if path.suffix.lower() == ".xlsx":
        return read_xlsx_records(path)
    raise ValueError("僅支援 .docx 或 .xlsx 名冊。")


def main() -> None:
    if len(sys.argv) != 2:
        raise SystemExit("請提供 Word 或 Excel 名冊路徑。")

    source = Path(sys.argv[1]).resolve()
    if not source.exists():
        raise SystemExit(f"找不到檔案：{source}")

    records = read_records(source)
    names = {record["name"] for record in records}
    output = Path(__file__).resolve().parents[1] / "staff-data.js"
    payload = json.dumps(records, ensure_ascii=False, indent=2)
    output.write_text(
        f"// 由 scripts/build_data.py 從正式名冊 {source.name} 產生，請勿手動編輯。\n"
        f"window.STAFF_DATA = {payload};\n",
        encoding="utf-8",
    )
    print(f"已產生 {output}：{len(records)} 筆任務、{len(names)} 位人員。")


if __name__ == "__main__":
    main()
