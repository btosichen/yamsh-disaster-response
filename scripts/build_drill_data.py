"""Build drill-data.js from the official matrix-style drill DOCX."""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


ROLE_LABELS = ["指揮官／發言人", "通報組", "避難引導組", "搶救組", "安全防護組", "緊急救護組", "老師", "學生"]
LEADERSHIP = [
    {"role": "指揮官", "name": "洪金英", "title": "校長"},
    {"role": "指揮官代理人", "name": "陳佑昌", "title": "教務主任"},
    {"role": "發言人", "name": "簡世煜", "title": "學務主任"},
]
SCENARIO_OVERRIDE = "115年09月21日9時21分（第二節上課）時段，臺北市發生震度（4級）地震，地震持續　30　秒（造成：▓電力及□通訊中斷，▓產生 求真樓設備組失火複合性災害），震動停止後，學校進行避難疏散。"
# Use a representative column near the centre of each merged role header.  Matrix
# body cells sometimes cross a header boundary, so using the left edge can pick
# up the preceding role's text.
FIRST_ROLE_STARTS = [7, 9, 12, 15, 18, 21, 23, 24]
SECOND_ROLE_STARTS = [8, 10, 13, 16, 19, 21, 23, 24]
TEAM_PROXIES = {
    "通報組": "楊麗雯",
    "避難引導組": "周峻民",
    "搶救組": "林承恩",
    "安全防護組": "王為倩",
    "緊急救護組": "林淳雅",
}


def clean(text: str) -> str:
    return text.strip().replace("\r", "")


def at(row, index: int) -> str:
    return clean(row.cells[index].text)


def unique_ranges(row) -> list[dict[str, object]]:
    ranges: list[dict[str, object]] = []
    previous = None
    for grid_index, cell in enumerate(row.cells):
        key = id(cell._tc)
        if key == previous:
            ranges[-1]["end"] = grid_index
        else:
            ranges.append({"start": grid_index, "end": grid_index, "text": clean(cell.text)})
            previous = key
    return ranges


def roster_team_summaries(path: Path) -> dict[str, str]:
    document = Document(path)
    groups: dict[str, dict[str, list[str]]] = {}
    for row in document.tables[0].rows[1:]:
        raw_group = re.sub(r"\s+", "", row.cells[0].text)
        if raw_group in {"指揮官", "指揮官代理人", "發言人"}:
            continue
        group = raw_group.split("（", 1)[0]
        role = re.sub(r"\s+", "", row.cells[1].text)
        names = [p.text.strip() for p in row.cells[2].paragraphs if p.text.strip() and not p.text.strip().startswith("*")]
        groups.setdefault(group, {})[role] = names
    summaries = {}
    for group, rows in groups.items():
        leaders = "、".join(rows.get("組長", []))
        members = rows.get("組員", [])
        parts = [f"組長：{leaders}"]
        proxy = TEAM_PROXIES.get(group)
        if proxy and proxy in members:
            parts.append(f"代理：{proxy}")
        parts.append(f"組員{len(members)}人")
        if group == "搶救組":
            parts.append("*高中專任（含外師）另列附表")
        elif group == "避難引導組":
            parts.append("*國高中導師另列附表")
        elif group == "安全防護組":
            parts.append("*國中專任另列附表")
        summaries[group] = "\n".join(parts)
    return summaries


def regular_step(table, row_no: int, layout: str) -> dict[str, object]:
    row = table.rows[row_no - 1]
    if layout == "first":
        scenario = at(row, 3)
        role_starts = FIRST_ROLE_STARTS
    else:
        scenario_parts = [at(row, 3), at(row, 4)]
        scenario = "\n".join(dict.fromkeys(part for part in scenario_parts if part))
        role_starts = SECOND_ROLE_STARTS
    actions = [
        {"role": role, "text": at(row, start)}
        for role, start in zip(ROLE_LABELS, role_starts)
    ]
    # In two medical-response rows, the command cell begins one grid column
    # before the second-half header boundary. Preserve it as the commander text.
    if layout == "second" and at(row, 6).startswith("指揮官"):
        actions[0]["text"] = at(row, 6)
    return {
        "row": row_no,
        "stage": at(row, 0),
        "time": at(row, 2),
        "scenario": scenario,
        "actions": actions,
    }


def note_step(table, row_no: int, title: str, starts: list[int]) -> dict[str, object]:
    row = table.rows[row_no - 1]
    return {
        "row": row_no,
        "stage": at(row, 0) if row_no not in {20, 33} else title,
        "time": at(row, 2) if row_no in {14, 16, 31} else "",
        "scenario": title,
        "notes": [at(row, start) for start in starts if at(row, start)],
    }


def build(source: Path, roster: Path | None = None) -> dict[str, object]:
    document = Document(source)
    table = document.tables[0]
    settings = []
    for row_no in range(2, 9):
        row = table.rows[row_no - 1]
        settings.append({"label": at(row, 1), "text": at(row, 5)})

    team_header = table.rows[9]
    team_values = table.rows[10]
    team_starts = FIRST_ROLE_STARTS
    teams = []
    roster_summaries = roster_team_summaries(roster) if roster else {}
    for start in team_starts:
        role = at(team_header, start)
        members = at(team_values, start)
        role_key = re.sub(r"\s+", "", role)
        if role_key in roster_summaries:
            members = roster_summaries[role_key]
        elif role == "通報組" and "丁肆山" not in members:
            members = f"{members}\n名冊更新：新增丁肆山（保全），通報組共10名組員"
        teams.append({"role": role, "members": "" if members == role else members})

    timeline = [
        regular_step(table, 13, "first"),
        note_step(table, 14, "判斷原則", [6]),
        regular_step(table, 15, "first"),
        note_step(table, 16, "全校廣播", [6]),
        regular_step(table, 17, "first"),
        regular_step(table, 18, "first"),
        regular_step(table, 19, "first"),
        note_step(table, 20, "情境組合提示", [0]),
        regular_step(table, 21, "second"),
        regular_step(table, 22, "second"),
        note_step(table, 23, "受傷類型與應變支援", [2, 8]),
        note_step(table, 24, "受困類型與回報要領", [2, 8]),
        regular_step(table, 27, "second"),
        regular_step(table, 28, "second"),
        regular_step(table, 29, "second"),
        regular_step(table, 30, "second"),
        note_step(table, 31, "裝備檢整與歸位", [7]),
        note_step(table, 32, "檢討會", [2]),
        note_step(table, 33, "腳本調製提醒", [0]),
    ]

    reference = document.tables[1]
    reference_headers = [clean(cell.text) for cell in reference.rows[1].cells]
    reference_rows = [
        [clean(cell.text) for cell in row.cells]
        for row in reference.rows[2:]
    ]
    title = at(table.rows[0], 0)
    match = re.search(r"更新日期：\s*([0-9.]+)", title)
    return {
        "title": title.split("更新日期：")[0].strip(),
        "updated": match.group(1) if match else "",
        "unitNotes": [clean(p.text) for p in document.paragraphs if p.text.strip()],
        "settingsTitle": at(table.rows[1], 0),
        "settings": settings,
        "scenarioLabel": at(table.rows[8], 0),
        "scenario": SCENARIO_OVERRIDE,
        "leadership": LEADERSHIP,
        "matrixHeaders": list(
            dict.fromkeys(
                at(row, index)
                for row, indexes in ((table.rows[9], range(25)), (table.rows[24], range(25)))
                for index in indexes
                if at(row, index)
            )
        ),
        "teams": teams,
        "oralReportNote": at(table.rows[11], 0),
        "timeline": timeline,
        "reference": {
            "title": at(reference.rows[0], 0),
            "headers": reference_headers,
            "rows": reference_rows,
        },
        "sourceFile": source.name,
    }


def main() -> None:
    if len(sys.argv) not in {2, 3}:
        raise SystemExit("請提供防災演練腳本 DOCX 路徑，以及選填的最新分組表 DOCX。")
    source = Path(sys.argv[1]).resolve()
    if not source.exists():
        raise SystemExit(f"找不到檔案：{source}")
    roster = Path(sys.argv[2]).resolve() if len(sys.argv) == 3 else None
    data = build(source, roster)
    output = Path(__file__).resolve().parents[1] / "drill-data.js"
    output.write_text(
        f"// 由 scripts/build_drill_data.py 從 {source.name} 產生，請勿手動編輯。\n"
        f"window.DRILL_DATA = {json.dumps(data, ensure_ascii=False, indent=2)};\n",
        encoding="utf-8",
    )
    print(f"已產生 {output}：{len(data['settings'])} 項設定、{len(data['timeline'])} 個演練步驟。")


if __name__ == "__main__":
    main()
