"""Update the emergency roster from the 115-09-07 staff workbook.

The existing DOCX remains the layout and duty-text authority.  The workbook is
the personnel/title authority.  Every person receives exactly one assignment.
"""

from __future__ import annotations

import copy
import re
import shutil
import sys
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from openpyxl import load_workbook


DATE_LABEL = "115.09.07新版"
EXTERNAL_SUPPORT = {
    "丁肆山": {"name": "丁肆山", "title": "保全", "group": "通報組", "role": "組員", "table": 1}
}
NEW_MAIN_ASSIGNMENTS = {
    "張安莛": ("緊急救護組", "組長"),
    "周軒如": ("通報組", "組員"),
    "李彥薇": ("通報組", "組員"),
    "李翊誠": ("通報組", "組員"),
    "陳芊穎": ("通報組", "組員"),
    "尹德融": ("搶救組", "組員"),
    "翁竹毅": ("搶救組", "組員"),
    "廖雪利": ("搶救組", "組員"),
    "洪志善": ("搶救組", "組員"),
    "賴致宇": ("搶救組", "組員"),
    "黃瑋琪": ("搶救組", "組員"),
    "林士恒": ("搶救組", "組員"),
    "范瑋婕": ("安全防護組", "組員"),
    "林婷婷": ("搶救組", "組員"),
}
MAIN_GROUP_ORDER = ["搶救組", "通報組", "避難引導組", "安全防護組", "緊急救護組"]
ANNEX_GROUPS = {2: "搶救組", 3: "避難引導組", 4: "安全防護組"}
GROUP_TO_TABLE = {"搶救組": 2, "避難引導組": 3, "安全防護組": 4}


def display(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def key(value) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip()


def source_title(title, unit="") -> str:
    title = key(title)
    unit = key(unit).replace("考分", "").strip()
    if title in {"主任", "組長", "組員", "助理員", "幹事", "管理員", "技佐", "書記"} and unit:
        return f"{unit}{title}"
    if title == "約僱人員" and unit:
        return f"{unit}約僱人員"
    return title


def teacher_title(subject, default="教師") -> str:
    value = key(subject)
    value = re.sub(r"(?:懸代|進修|育嬰)", "", value)
    value = re.sub(r"[（(]無教證[）)]", "", value)
    value = re.sub(r"\d{3}導師", "", value)
    value = value.strip("-_—")
    if not value:
        return default
    return value if value.endswith("教師") else f"{value}教師"


def add(records, name, title, section, detail="", active=True):
    name = display(name)
    if not name:
        return
    records.setdefault(key(name), {
        "name": name,
        "title": source_title(title, detail),
        "section": section,
        "detail": display(detail),
        "active": active,
    })


def read_staff(path: Path):
    ws = load_workbook(path, data_only=True, read_only=True).active
    records = OrderedDict()
    for row in range(3, 53):
        _, title, name, detail = (ws.cell(row, col).value for col in range(1, 5))
        if not name:
            continue
        section, active = "教師兼行政", True
        if 27 <= row <= 29:
            section = "運動教練"
        elif row == 31:
            section, title = "約僱人員", "控障教師約僱人員"
        elif 35 <= row <= 38:
            section = "編制臨時人員"
        elif 40 <= row <= 44:
            section, active = "留停或支援", False
        elif 47 <= row <= 49:
            section = "外籍教師"
        elif 50 <= row <= 51:
            section = "專案行政助理"
        elif row == 52:
            section = "救生員"
        add(records, name, title, section, detail, active)

    for row in range(4, 54):
        _, class_no, name, subject = (ws.cell(row, col).value for col in range(5, 9))
        if not name:
            continue
        subject_text = display(subject)
        if row <= 44:
            add(records, name, f"{key(class_no)}導師", "國高中導師", subject)
        elif row == 46:
            add(records, name, "高中特教導師", "國高中導師", subject)
        else:
            match = re.search(r"(\d{3})\s*導師", subject_text)
            if match:
                add(records, name, f"{match.group(1)}導師", "國高中導師", subject)
            else:
                add(records, name, teacher_title(subject_text, "高中教師"), "高中專任", subject)

    for row in range(3, 50):
        _, _, name, subject = (ws.cell(row, col).value for col in range(9, 13))
        if name:
            add(records, name, teacher_title(subject, "高中教師"), "高中專任", subject)

    for row in range(4, 49):
        _, class_no, name, subject = (ws.cell(row, col).value for col in range(13, 17))
        if not name:
            continue
        if row <= 24:
            add(records, name, f"{key(class_no)}導師", "國高中導師", subject)
        else:
            add(records, name, teacher_title(subject, "國中教師"), "國中專任", subject)

    for row in range(3, 42):
        _, _, title, name, unit = (ws.cell(row, col).value for col in range(17, 22))
        if name:
            add(records, name, title, "職員職工", unit)
    return OrderedDict((k, v) for k, v in records.items() if v["active"])


def cell_items(cell):
    return [display(p.text) for p in cell.paragraphs if display(p.text)]


def parse_current(document):
    assignments = {}
    for table_index, table in enumerate(document.tables, 1):
        for row in table.rows[1:]:
            names, titles = cell_items(row.cells[2]), cell_items(row.cells[3])
            if len(names) != len(titles):
                raise ValueError(f"姓名與職稱未對齊：表{table_index} {row.cells[0].text}")
            raw_group = key(row.cells[0].text)
            if table_index == 1:
                group = raw_group if raw_group in {"指揮官", "指揮官代理人", "發言人"} else raw_group.split("（", 1)[0]
            else:
                group = ANNEX_GROUPS[table_index]
            role = key(row.cells[1].text)
            for name, title in zip(names, titles):
                if name.startswith("*"):
                    continue
                assignments[key(name)] = {"name": name, "title": title, "group": group, "role": role, "table": table_index}
    return assignments


def choose_assignment(record, current):
    name_key = key(record["name"])
    if name_key in {key(n) for n in NEW_MAIN_ASSIGNMENTS}:
        group, role = NEW_MAIN_ASSIGNMENTS[next(n for n in NEW_MAIN_ASSIGNMENTS if key(n) == name_key)]
        return group, role, 1
    # The current main table contains deliberate specialist assignments (for
    # example counsellors in the medical team). Preserve those assignments;
    # only generic annex placement follows the newest teaching role/class.
    if current and current["table"] == 1:
        return current["group"], current["role"], 1
    section = record["section"]
    if section == "國高中導師":
        return "避難引導組", "組員", 3
    if section == "高中專任" or section == "外籍教師":
        if "國中" in record["title"]:
            return "安全防護組", "組員", 4
        return "搶救組", "組員", 2
    if section == "國中專任":
        return "安全防護組", "組員", 4
    if current:
        return current["group"], current["role"], 1
    raise ValueError(f"沒有分組規則：{record}")


def clear_and_fill(cell, lines, template_cell):
    template_p = template_cell.paragraphs[0]
    template_ppr = copy.deepcopy(template_p._p.pPr)
    template_rpr = copy.deepcopy(template_p.runs[0]._r.rPr) if template_p.runs else None
    for paragraph in list(cell.paragraphs):
        cell._tc.remove(paragraph._p)
    for text in lines:
        paragraph = cell.add_paragraph()
        if template_ppr is not None:
            paragraph._p.insert(0, copy.deepcopy(template_ppr))
        run = paragraph.add_run(text)
        if template_rpr is not None:
            run._r.insert(0, copy.deepcopy(template_rpr))
    if not lines:
        cell.add_paragraph()


def replace_first_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def compact_member_cell(cell, size: float):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 0.9
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)


def prevent_row_split(row):
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def build(staff_path: Path, source_docx: Path, output_docx: Path):
    document = Document(source_docx)
    staff = read_staff(staff_path)
    current = parse_current(document)
    assignments = OrderedDict()
    renamed_keys = {key("李桂鳳"): key("李桂鳯")}

    for staff_key, record in staff.items():
        lookup = current.get(staff_key)
        if lookup is None:
            old_key = next((old for old, new in renamed_keys.items() if new == staff_key), None)
            lookup = current.get(old_key) if old_key else None
        group, role, table = choose_assignment(record, lookup)
        assignments[staff_key] = {**record, "group": group, "role": role, "table": table}
    for name, record in EXTERNAL_SUPPORT.items():
        assignments[key(name)] = dict(record)

    # Preserve command roles and explicit group leaders, but use current staff titles.
    leadership = {
        "洪金英": ("指揮官", "指揮官"),
        "陳佑昌": ("指揮官代理人", "指揮官代理人"),
        "簡世煜": ("發言人", "發言人"),
        "陳政川": ("搶救組", "組長"),
        "吳家瑋": ("通報組", "組長"),
        "王聖淵": ("避難引導組", "組長"),
        "王秋堯": ("安全防護組", "組長"),
        "張安莛": ("緊急救護組", "組長"),
        "林巧芳": ("緊急救護組", "組長"),
    }
    for name, (group, role) in leadership.items():
        record = assignments[key(name)]
        record.update(group=group, role=role, table=1)

    main = {group: {"組長": [], "組員": []} for group in MAIN_GROUP_ORDER}
    command = {}
    annex = {2: [], 3: [], 4: []}
    for record in assignments.values():
        if record["table"] == 1:
            if record["group"] in {"指揮官", "指揮官代理人", "發言人"}:
                command[record["group"]] = record
            else:
                main[record["group"]][record["role"]].append(record)
        else:
            annex[record["table"]].append(record)

    # Retain the original main-table row structure.
    main_rows = {
        "指揮官": document.tables[0].rows[1],
        "指揮官代理人": document.tables[0].rows[2],
        "發言人": document.tables[0].rows[3],
        "搶救組組長": document.tables[0].rows[4], "搶救組組員": document.tables[0].rows[5],
        "通報組組長": document.tables[0].rows[6], "通報組組員": document.tables[0].rows[7],
        "避難引導組組長": document.tables[0].rows[8], "避難引導組組員": document.tables[0].rows[9],
        "安全防護組組長": document.tables[0].rows[10], "安全防護組組員": document.tables[0].rows[11],
        "緊急救護組組長": document.tables[0].rows[12], "緊急救護組組員": document.tables[0].rows[13],
    }
    for label, record in command.items():
        row = main_rows[label]
        clear_and_fill(row.cells[2], [record["name"]], row.cells[2])
        clear_and_fill(row.cells[3], [record["title"]], row.cells[3])

    for group in MAIN_GROUP_ORDER:
        for role in ("組長", "組員"):
            row = main_rows[f"{group}{role}"]
            members = main[group][role]
            names = [member["name"] for member in members]
            titles = [member["title"] for member in members]
            clear_and_fill(row.cells[2], names, row.cells[2])
            clear_and_fill(row.cells[3], titles, row.cells[3])
            if role == "組員":
                compact_member_cell(row.cells[2], 12)
                compact_member_cell(row.cells[3], 11)

    for table_index in (2, 3, 4):
        row = document.tables[table_index - 1].rows[1]
        members = annex[table_index]
        clear_and_fill(row.cells[2], [m["name"] for m in members], row.cells[2])
        clear_and_fill(row.cells[3], [m["title"] for m in members], row.cells[3])

    for paragraph in document.paragraphs:
        if paragraph.text.startswith("北市立陽明高級中學"):
            replace_first_text(paragraph, paragraph.text.replace("北市立", "臺北市立", 1))
        if key(paragraph.text) == "115.08.21新版":
            replace_first_text(paragraph, DATE_LABEL)
        if paragraph.text.startswith("三、國中專任教師"):
            replace_first_text(paragraph, "三、國中專任教師（安全防護組／安全防護班）")

    for row in document.tables[0].rows[1:]:
        prevent_row_split(row)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)
    print(f"完成：{output_docx}")
    print(f"共 {len(assignments)} 人（教職員與支援 {len(staff)} 人，保全 1 人），每人 1 項任務")
    print("主表：" + "、".join(f"{g}{sum(map(len, main[g].values()))}人" for g in MAIN_GROUP_ORDER))
    print("附表：" + "、".join(f"表{i} {len(annex[i])}人" for i in (2, 3, 4)))


if __name__ == "__main__":
    if len(sys.argv) != 4:
        raise SystemExit("usage: update_roster_1150907.py STAFF.xlsx SOURCE.docx OUTPUT.docx")
    build(*(Path(arg).resolve() for arg in sys.argv[1:]))
